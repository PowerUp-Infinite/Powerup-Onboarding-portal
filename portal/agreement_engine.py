"""
agreement_engine.py — Generate client agreement documents (DOCX).

Two agreement types:
  - Elite (PowerUp Infinite): Annexure A only — name + date + fee slabs
  - Non-Elite (UW IA Client Agreement): Full agreement — name, email, phone + fee slabs

Fee slabs:
  - Default: 3 slabs (0.75%, 0.625%, 0.5%) already in the template
  - Custom: user-supplied slabs (1–N rows), replaces the fee table entirely

Templates are downloaded from Google Drive at runtime.
"""

import os
import copy
import tempfile
from datetime import date
from io import BytesIO

from docx import Document
from docx.oxml.ns import qn
from googleapiclient.http import MediaIoBaseDownload

from google_auth import get_drive_service
import config

# ── Template download cache ───────────────────────────────────
_CACHE_DIR = tempfile.mkdtemp(prefix="agreement_")


def _download_template(file_id: str, filename: str) -> str:
    """Download a DOCX template from Drive (cached per session)."""
    path = os.path.join(_CACHE_DIR, filename)
    if os.path.exists(path):
        return path
    svc = get_drive_service()
    request = svc.files().get_media(fileId=file_id, supportsAllDrives=True)
    buf = BytesIO()
    downloader = MediaIoBaseDownload(buf, request)
    done = False
    while not done:
        _, done = downloader.next_chunk()
    with open(path, "wb") as f:
        f.write(buf.getvalue())
    return path


# ── Default fee slabs ─────────────────────────────────────────
DEFAULT_SLABS = [
    {"fee": "0.75% p.a.", "aua": "less than 50 lakhs"},
    {"fee": "0.625% p.a.", "aua": "\u20b950 lakhs to \u20b91 crore"},
    {"fee": "0.5% p.a.", "aua": "\u20b91 crore and above"},
]


# ── Table helpers ─────────────────────────────────────────────

def _clear_cell(cell):
    """Remove all content from a table cell while preserving formatting."""
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ""
    # Keep exactly one paragraph
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]._element
        p.getparent().remove(p)


def _set_cell_text(cell, text: str):
    """Set cell text, preserving the formatting of the first run."""
    _clear_cell(cell)
    if cell.paragraphs[0].runs:
        cell.paragraphs[0].runs[0].text = text
    else:
        cell.paragraphs[0].text = text


def _delete_table_row(table, row_idx: int):
    """Delete a row from a table by index."""
    row = table.rows[row_idx]
    tbl = table._tbl
    tbl.remove(row._tr)


def _rebuild_fee_table(table, slabs: list[dict]):
    """
    Replace the fee slab rows in the table.
    Row 0 is the header; rows 1+ are slab rows.
    """
    # Delete existing data rows (bottom-up to preserve indices)
    for i in range(len(table.rows) - 1, 0, -1):
        _delete_table_row(table, i)

    # Add new slab rows
    for idx, slab in enumerate(slabs):
        # Clone the header row's XML to get consistent formatting
        new_tr = copy.deepcopy(table.rows[0]._tr)
        table._tbl.append(new_tr)
        # Now set cell values in the newly added row
        new_row = table.rows[-1]
        _set_cell_text(new_row.cells[0], str(idx + 1))
        _set_cell_text(new_row.cells[1], slab["fee"])
        _set_cell_text(new_row.cells[2], slab["aua"])


# ── Paragraph helpers ─────────────────────────────────────────

def _set_paragraph_field(paragraph, prefix: str, value: str):
    """
    Replace text after a prefix in a paragraph.
    E.g., paragraph "Name: Old Name" → "Name: New Name"
    Handles multi-run paragraphs by clearing all runs after the prefix run
    and setting the value in the second run.
    """
    runs = paragraph.runs
    if not runs:
        return

    # Find which run contains the prefix
    full_text = paragraph.text
    if prefix not in full_text:
        return

    # Strategy: set the full text as "prefix value" in the first run,
    # clear all other runs
    first_run = runs[0]
    first_run.text = f"{prefix}{value}"
    for r in runs[1:]:
        r.text = ""


# ── Elite agreement ───────────────────────────────────────────

def generate_elite(
    client_name: str,
    agreement_date: date | None = None,
    custom_slabs: list[dict] | None = None,
) -> tuple[BytesIO, str]:
    """
    Generate an Elite (PowerUp Infinite) agreement.

    Args:
        client_name: Full client name
        agreement_date: Date for the agreement (defaults to today)
        custom_slabs: If provided, replaces the default fee slabs.
                      Each dict: {"fee": "0.75% p.a.", "aua": "less than 50 lakhs"}

    Returns:
        (docx_bytes: BytesIO, filename: str)
    """
    if agreement_date is None:
        agreement_date = date.today()

    template_path = _download_template(
        config.AGREEMENT_ELITE_TEMPLATE_ID,
        "elite_template.docx",
    )
    doc = Document(template_path)

    # -- Fill name (P39: "Name: ...")
    for p in doc.paragraphs:
        if p.text.startswith("Name:"):
            _set_paragraph_field(p, "Name: ", client_name)
            break

    # -- Fill date (P40: "Date: ...")
    date_str = agreement_date.strftime("%d-%m-%Y")
    for p in doc.paragraphs:
        if p.text.startswith("Date:"):
            _set_paragraph_field(p, "Date: ", date_str)
            break

    # -- Fee slabs (Table 0)
    if custom_slabs:
        _rebuild_fee_table(doc.tables[0], custom_slabs)

    # -- Save to BytesIO
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"PowerUp Infinite Agreement - {client_name}.docx"
    return buf, filename


# ── Non-elite agreement ───────────────────────────────────────

def generate_nonelite(
    client_name: str,
    email: str = "",
    phone: str = "",
    custom_slabs: list[dict] | None = None,
) -> tuple[BytesIO, str]:
    """
    Generate a Non-Elite (UW IA Client Agreement with Annex A) agreement.

    Args:
        client_name: Full client name
        email: Client email (if "NA" or empty, the Email row is deleted)
        phone: Client phone (if "NA" or empty, the Phone row is deleted)
        custom_slabs: If provided, replaces the default fee slabs.

    Returns:
        (docx_bytes: BytesIO, filename: str)
    """
    template_path = _download_template(
        config.AGREEMENT_NONELITE_TEMPLATE_ID,
        "nonelite_template.docx",
    )
    doc = Document(template_path)

    # -- Fee slabs (Table 0)
    if custom_slabs:
        _rebuild_fee_table(doc.tables[0], custom_slabs)

    # -- Schedule I (Table 1): client details
    schedule = doc.tables[1]

    # Row 0: Name
    _set_cell_text(schedule.rows[0].cells[1], client_name)

    # Track rows to delete (bottom-up later)
    rows_to_delete = []

    # Row 2: Email
    if _is_na(email):
        rows_to_delete.append(2)
    else:
        _set_cell_text(schedule.rows[2].cells[1], email)

    # Row 3: Phone
    if _is_na(phone):
        rows_to_delete.append(3)
    else:
        _set_cell_text(schedule.rows[3].cells[1], phone)

    # Delete flagged rows bottom-up
    for ri in sorted(rows_to_delete, reverse=True):
        _delete_table_row(schedule, ri)

    # -- Save to BytesIO
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"UW IA Client Agreement - {client_name}.docx"
    return buf, filename


def _is_na(val: str) -> bool:
    """Check if a value is effectively 'not applicable'."""
    if not val:
        return True
    return val.strip().upper() in ("NA", "N/A", "-", "")
